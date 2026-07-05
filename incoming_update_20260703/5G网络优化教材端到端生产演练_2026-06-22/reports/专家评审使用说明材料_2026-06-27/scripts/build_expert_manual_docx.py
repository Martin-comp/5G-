from __future__ import annotations

import json
import math
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/ny/Downloads/5G网络优化教材端到端生产演练_2026-06-22")
PKG = ROOT / "reports/专家评审使用说明材料_2026-06-27"
FIG_DIR = PKG / "figures"
SCREEN_DIR = FIG_DIR / "screenshots_v0_2"
DOC_IMG_DIR = FIG_DIR / "doc_images_v0_2"
DIAGRAM_DIR = FIG_DIR / "diagrams_v0_2"
OUT_DOCX = PKG / "5G网络优化教材数字教材预览版_专家评审使用说明V0.2.docx"
SOURCE_MD = PKG / "5G网络优化教材数字教材预览版_专家评审使用说明V0.2.md"

PREVIEW_LINK = ROOT / "samples/digital_textbook_overall_prototype_overall_0_5_clean_preview/index.html"
TASK_LINK = ROOT / "samples/digital_textbook_overall_prototype_overall_0_5_clean_preview/task_workbench_3_5a1_two_period_sample/index.html"

ACCENT = "166A6A"
ACCENT_DARK = "0E4F4F"
MUTED = "566575"
LINE = "D8E2EA"
FILL = "F3F8F8"
SOFT_GOLD = "FFF5E6"
GOLD = "9C5B1A"


def ensure_dirs() -> None:
    DOC_IMG_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
    ]
    for item in candidates:
        try:
            return ImageFont.truetype(item, size)
        except Exception:
            pass
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in str(text).split("\n"):
        if not raw:
            lines.append("")
            continue
        line = ""
        for char in raw:
            trial = line + char
            if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
                line = trial
            else:
                if line:
                    lines.append(line)
                line = char
        if line:
            lines.append(line)
    return lines


def draw_wrapped(draw, xy, text, fnt, fill, max_width, line_gap=8):
    x, y = xy
    for line in wrap_text(draw, text, fnt, max_width):
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def rounded_box(draw, box, fill, outline, width=2, radius=22):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#6B8FA4", width=5):
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    angle = math.atan2(ey - sy, ex - sx)
    size = 16
    p1 = (ex - size * math.cos(angle - math.pi / 6), ey - size * math.sin(angle - math.pi / 6))
    p2 = (ex - size * math.cos(angle + math.pi / 6), ey - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=fill)


def make_diagram_architecture() -> Path:
    out = DIAGRAM_DIR / "01_平台总体架构图.png"
    img = Image.new("RGB", (1800, 1000), "#F6FAFC")
    d = ImageDraw.Draw(img)
    title = font(42, True)
    h = font(28, True)
    b = font(22)
    small = font(19)
    d.text((70, 55), "数字教材预览版总体架构", font=title, fill=f"#{ACCENT_DARK}")
    d.text((72, 115), "从输入材料到课程能力图谱，再到整书预览、任务学习、平台支持和专家反馈的关系", font=b, fill="#4F5E69")

    boxes = [
        ((70, 220, 360, 450), "输入材料", "传统电子教材\n课程标准/教学大纲\n互联网可信补充\n已有图片/表格/视频"),
        ((440, 220, 760, 450), "课程能力图谱主数据", "课程主链\n项目任务\n能力节点\n资源挂接\n审核状态"),
        ((850, 180, 1230, 500), "整书预览层", "课程\n项目\n图谱\n教师\n平台支持辅助入口"),
        ((1320, 220, 1710, 450), "任务级深样章", "P4-T2结果验证\n学生学习\n任务组织\n任务资源"),
        ((850, 610, 1230, 850), "专家评审与迭代", "教学可用性\n专业正确性\n图谱清晰度\n资源适配度\n平台交付建议"),
    ]
    fills = ["#FFFFFF", "#EAF7F5", "#FFFFFF", "#FFF7EA", "#FFFFFF"]
    for index, (box, label, desc) in enumerate(boxes):
        rounded_box(d, box, fills[index], "#B8CCD8", 3)
        x1, y1, x2, _ = box
        d.text((x1 + 28, y1 + 28), label, font=h, fill=f"#{ACCENT_DARK}" if index in [1, 2] else "#1F2D3A")
        draw_wrapped(d, (x1 + 28, y1 + 84), desc, b if index != 4 else small, "#33414D", x2 - x1 - 56, 9)

    arrow(d, (360, 335), (440, 335))
    arrow(d, (760, 335), (850, 335))
    arrow(d, (1230, 335), (1320, 335))
    arrow(d, (1040, 500), (1040, 610))
    arrow(d, (850, 720), (760, 430))
    d.text((625, 620), "反馈回写图谱和内容生产规则", font=small, fill=f"#{GOLD}")
    img.save(out)
    return out


def make_diagram_graph_model() -> Path:
    out = DIAGRAM_DIR / "02_课程能力图谱关系模型.png"
    img = Image.new("RGB", (1800, 950), "#FBFCFD")
    d = ImageDraw.Draw(img)
    title = font(42, True)
    h = font(26, True)
    b = font(21)
    d.text((70, 55), "课程能力图谱关系模型", font=title, fill=f"#{ACCENT_DARK}")
    d.text((72, 116), "图谱不是目录，而是连接项目、任务、资源、活动和评价产出的暗线", font=b, fill="#4F5E69")

    levels = [
        ("第一层：课程主链", "信息采集 → 网络测试 → 信息管理 → 优化实施 → 结果验证 → 性能提升 → 信令分析"),
        ("第二层：项目路径", "P2网络测试路径与P4端到端优化路径形成证据承接"),
        ("第三层：能力节点", "每个节点表示一个可学习、可训练、可评价的具体动作"),
        ("第四层：资源卡片", "学习单、表格、互动、页面、图片或视频服务具体节点训练"),
        ("评价产出", "学生完成可检查的学习产出，教师据此讲评和验收"),
    ]
    y = 210
    prev = None
    for i, (label, desc) in enumerate(levels):
        box = (160, y, 1640, y + 115)
        rounded_box(d, box, "#EAF7F5" if i in [0, 2] else "#FFFFFF", "#B8CCD8", 3)
        d.text((195, y + 22), label, font=h, fill=f"#{ACCENT_DARK}")
        draw_wrapped(d, (500, y + 28), desc, b, "#33414D", 1050)
        if prev:
            arrow(d, (900, prev + 115), (900, y - 10), width=4)
        prev = y
        y += 145
    img.save(out)
    return out


def make_diagram_review_route() -> Path:
    out = DIAGRAM_DIR / "03_专家评审推荐路径.png"
    img = Image.new("RGB", (1800, 950), "#F6FAFC")
    d = ImageDraw.Draw(img)
    title = font(42, True)
    h = font(25, True)
    b = font(20)
    d.text((70, 55), "专家评审推荐路径", font=title, fill=f"#{ACCENT_DARK}")
    d.text((72, 116), "建议专家先看整体，再看图谱，最后下钻到P4-T2任务样章和平台支持，并按清单提出改进建议", font=b, fill="#4F5E69")

    steps = [
        ("1 课程首页", "理解整书入口和项目链"),
        ("2 项目页", "看P2/P4重点路径"),
        ("3 图谱页", "判断能力图谱是否清楚"),
        ("4 P4-T2任务", "模拟学生学习闭环"),
        ("5 教师页", "评估AI辅助组织与复核"),
        ("6 平台支持", "评估资源治理和平台交付"),
    ]
    x_positions = [80, 360, 640, 920, 1200, 1480]
    y = 340
    for i, (label, desc) in enumerate(steps):
        x = x_positions[i]
        rounded_box(d, (x, y, x + 230, y + 185), "#FFFFFF", "#B8CCD8", 3)
        d.text((x + 24, y + 26), label, font=h, fill=f"#{ACCENT_DARK}")
        draw_wrapped(d, (x + 24, y + 80), desc, b, "#33414D", 178)
        if i < len(steps) - 1:
            arrow(d, (x + 230, y + 92), (x_positions[i + 1] - 12, y + 92), width=4)
    rounded_box(d, (360, 650, 1440, 805), "#FFF7EA", "#E6C58E", 3)
    d.text((400, 690), "评审输出建议", font=h, fill=f"#{GOLD}")
    draw_wrapped(d, (620, 690), "按“问题位置、影响、修改建议、优先级、是否需专业复核”提交反馈，便于后续迭代。", b, "#3D3D3D", 740)
    img.save(out)
    return out


def crop(src_name: str, out_name: str, y0: int = 0, height: int | None = None) -> Path:
    src = SCREEN_DIR / src_name
    out = DOC_IMG_DIR / out_name
    im = Image.open(src).convert("RGB")
    w, h = im.size
    if height is None:
        crop_box = (0, y0, w, h)
    else:
        crop_box = (0, y0, w, min(h, y0 + height))
    im = im.crop(crop_box)
    # Light border on white pages so screenshots remain visible in Word.
    canvas = Image.new("RGB", (im.width + 24, im.height + 24), "#FFFFFF")
    canvas.paste(im, (12, 12))
    d = ImageDraw.Draw(canvas)
    d.rectangle((8, 8, canvas.width - 9, canvas.height - 9), outline="#D8E2EA", width=3)
    canvas.save(out, quality=92)
    return out


def prepare_doc_images() -> dict[str, Path]:
    images = {
        "course": crop("01_课程首页_整书入口.png", "01_课程首页_整书入口.png"),
        "project": crop("02_项目页_项目链与学习路径.png", "02_项目页_项目链与学习路径.png"),
        "project_p4": crop("03_项目页_项目四路径.png", "03_项目页_项目四路径.png"),
        "graph_top": crop("04_图谱页_课程能力图谱分层关系.png", "04_图谱页_课程能力图谱_上半部分.png", 0, 1220),
        "graph_resource": crop("05_图谱页_项目四能力节点与资源挂接.png", "05_图谱页_项目四资源挂接.png", 1080, 1250),
        "teacher": crop("06_教师页_整书教学组织入口.png", "06_教师页_整书教学组织入口.png"),
        "support_resources": crop("07_平台支持_任务资源映射.png", "07_平台支持_任务资源映射.png", 0, 1320),
        "support_material": crop("08_平台支持_素材子平台说明.png", "08_平台支持_素材子平台说明.png", 0, 1320),
        "support_delivery": crop("09_平台支持_交付说明.png", "09_平台支持_交付说明.png", 0, 1320),
        "task_student": crop("10_P4T2学生学习页_任务闭环.png", "10_P4T2学生学习页_任务闭环.png"),
        "task_teacher": crop("11_P4T2任务组织页_课堂组织.png", "11_P4T2任务组织页_课堂组织_上半部分.png", 0, 1320),
        "task_teacher_data": crop("11_P4T2任务组织页_课堂组织.png", "12_P4T2任务组织页_AI建议与复核.png", 1320, 1320),
        "task_resources": crop("12_P4T2任务资源页_资源转化.png", "13_P4T2任务资源页_资源转化.png", 0, 1320),
    }
    images["diagram_arch"] = make_diagram_architecture()
    images["diagram_graph"] = make_diagram_graph_model()
    images["diagram_route"] = make_diagram_review_route()
    return images


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, ACCENT_DARK, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "5G网络优化教材（高级）数字教材预览版 | 专家评审使用说明"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "供专家评审使用：请结合专业正确性、教学可用性、图谱清晰度和平台交付提出意见。"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=8.5, color=MUTED)


def add_para(doc: Document, text: str = "", style=None, bold=False, color=None, size=None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 12), color=ACCENT if level < 3 else ACCENT_DARK, bold=True)
    return p


def add_callout(doc: Document, title: str, body: str, fill: str = FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=11, color=ACCENT_DARK, bold=True)
    p.paragraph_format.space_after = Pt(4)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color="333333")
    p2.paragraph_format.space_after = Pt(2)
    return table


def add_kv_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2100, 7260])
    for idx, (k, v) in enumerate(rows):
        if idx > 0:
            table.add_row()
        c0, c1 = table.rows[idx].cells
        set_cell_shading(c0, "F2F4F7")
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(k)
        set_run_font(r0, size=10.5, color=ACCENT_DARK, bold=True)
        r1 = c1.paragraphs[0].add_run(v)
        set_run_font(r1, size=10.5, color="333333")
    return table


def add_caption(doc: Document, text: str):
    p = add_para(doc)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)
    return p


def add_image(doc: Document, path: Path, caption: str, width_in: float = 6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=10.8)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, size=10.8)


def add_review_table(doc: Document):
    headers = ["评审维度", "重点关注与建议反馈"]
    rows = [
        ("整体可理解性", "专家首次打开后，是否能理解课程、项目、图谱、任务和资源之间的关系。请指出卡住的位置，并说明是导航、文字、术语还是页面结构导致。"),
        ("教学可用性", "教师是否能据此组织课堂，学生是否能按任务完成学习产出。请给出课堂使用建议、学生可能不懂的位置和需要补充的支架。"),
        ("专业正确性", "5G网优指标、判断口径、案例数据、结论表达是否可靠。请标注错误或待复核口径，并给出正确表达或可采用标准。"),
        ("图谱支撑性", "课程能力图谱是否能解释前后承接、节点训练和资源挂接。请指出节点缺失、关系不清、资源挂接不合理或评价产出不明确的位置。"),
        ("资源转化价值", "图片、表格、互动、动画、小游戏等是否服务学习任务。请建议哪些资源应保留、重绘、替换、转成互动或删除。"),
        ("平台交付风险", "是否存在版权、数据安全、媒体重绘、平台接口和质量检测问题。请按高/中/低优先级说明风险和处理建议。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2200, 7160])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10.5, color=ACCENT_DARK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(text)
            set_run_font(run, size=10.0, color="333333")
    return table


def add_step_table(doc: Document, steps: list[str]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1300, 8060])
    headers = ["步骤", "操作"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10.5, color=ACCENT_DARK, bold=True)
    for index, step in enumerate(steps, start=1):
        cells = table.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        r0 = cells[0].paragraphs[0].add_run(str(index))
        set_run_font(r0, size=10.3, color=ACCENT_DARK, bold=True)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = cells[1].paragraphs[0].add_run(step)
        set_run_font(r1, size=10.3, color="333333")
    return table


def markdown_source() -> str:
    return f"""# 5G网络优化教材数字教材预览版专家评审使用说明 V0.2

本文件是同名 Word 文档的内容源，最终排版以 docx 为准。

入口：{PREVIEW_LINK}
任务样章：{TASK_LINK}

核心阅读路径：课程首页 -> 项目页 -> 图谱页 -> P4-T2任务 -> 教师页 -> 平台支持。
"""


def build_doc(images: dict[str, Path]) -> None:
    doc = Document()
    style_doc(doc)

    # Cover.
    p = add_para(doc, "专家评审使用说明", bold=True, color=ACCENT_DARK, size=12)
    p.paragraph_format.space_after = Pt(16)
    p = add_para(doc, "5G网络优化教材（高级）数字教材预览版", bold=True, color="111827", size=24)
    p.paragraph_format.space_after = Pt(4)
    add_para(doc, "面向专家评审的图文操作指南与改进建议收集说明", color=MUTED, size=13)
    add_para(doc, "版本：V0.2    日期：2026-06-29", color=MUTED, size=10)
    add_callout(
        doc,
        "阅读定位",
        "本说明用于帮助专家快速理解当前数字教材预览版的整体结构、课程能力图谱作用、任务级深样章呈现方式、教师/学生使用路径和资源交付逻辑。它不是教材定稿说明，也不替代专业复核、媒体审查或真实教学试用。",
    )
    add_kv_table(doc, [
        ("适用对象", "评审专家、职业院校教师、出版社与平台建设人员。"),
        ("整书预览入口", str(PREVIEW_LINK)),
        ("任务级深样章入口", str(TASK_LINK)),
        ("建议评审时间", "先用15分钟按本说明走完整体路径，再用30-60分钟重点评审P4-T2任务样章与课程能力图谱。"),
        ("反馈建议", "优先指出影响专家理解、教师教学、学生自学、专业正确性和平台交付的具体问题。"),
    ])
    doc.add_page_break()

    add_heading(doc, "一、专家应先了解的整体结论", 1)
    add_para(doc, "当前版本是一套“正式呈现观察版”的静态HTML预览，目的是让专家看到未来数字教材的大体阅读方式、学习路径、图谱组织方式和资源挂接逻辑。")
    add_bullets(doc, [
        "它已经把整书项目链、项目二网络测试路径、项目四端到端优化路径、P4-T2任务级深样章、AI辅助教师组织页和平台支持辅助入口放入同一个入口。",
        "它重点展示的是“数字教材如何组织学习”，不是展示已经完成全书所有内容生产。",
        "它保留了一个完整任务样章：P4-T2“5G网络优化结果验证”。专家可以通过这个任务观察学生学习闭环、教师课堂组织和资源转化方式。",
        "课程能力图谱当前作为暗线呈现，用于解释课程主链、项目路径、能力节点、资源卡片和评价产出之间的关系。",
    ])
    add_callout(
        doc,
        "当前边界",
        "请专家不要把本预览误认为正式出版稿。专业阈值、案例数据、媒体截图重绘、真实学生数据、出版社平台接口和质量检测仍需要后续专项复核。专家评审的重点是判断方向是否成立、结构是否清楚、教学是否可用、哪些地方必须修订。",
        SOFT_GOLD,
    )

    add_heading(doc, "二、平台总体架构", 1)
    add_para(doc, "平台当前采用“整书预览层 + 课程能力图谱 + 任务级深样章”的组合方式。专家阅读时可以把整书预览看作总入口，把图谱看作资源和任务组织的暗线，把P4-T2任务样章看作教材正文与学习活动的深度样本。")
    add_image(doc, images["diagram_arch"], "图1 平台总体架构：输入材料经课程能力图谱组织后进入整书预览、任务样章和专家反馈迭代。", 6.35)

    add_heading(doc, "三、推荐评审路径", 1)
    add_para(doc, "建议专家按“先整体、再图谱、再任务、再教师和资源、最后看交付”的顺序阅读。这样可以避免一开始就陷入某个页面细节，而忽略数字教材整体结构是否成立。")
    add_image(doc, images["diagram_route"], "图2 专家评审推荐路径：从整书入口逐步下钻到任务样章，再回看资源和交付。", 6.35)
    add_numbered(doc, [
        "打开整书预览入口，先确认课程项目链和重点学习路径。",
        "进入项目页，观察项目二与项目四之间的关系。",
        "进入图谱页，判断课程能力图谱是否能解释学习承接和资源挂接。",
        "进入P4-T2任务样章，模拟学生完成一次结果验证任务。",
        "查看教师页，判断AI预生成建议和教师审核确认是否能支持两课时课堂组织。",
        "查看平台支持，提出任务资源映射、素材子平台边界、平台接口和质量检测建议。",
    ])

    add_heading(doc, "四、按步骤查看整书预览版", 1)
    add_heading(doc, "步骤1：从课程首页理解整书入口", 2)
    add_para(doc, "课程首页用于快速说明本书覆盖哪些项目、哪些路径已经重点展开、哪里可以进入完整任务学习。专家重点看首屏是否能让一线教师和学生知道从哪里开始。")
    add_image(doc, images["course"], "图3 课程首页：展示整书项目链、重点路径和P4-T2任务入口。", 6.35)
    add_bullets(doc, [
        "顶部一级导航为课程、项目、图谱、教师；资源映射、素材子平台和交付说明进入“平台支持”辅助入口。",
        "课程页不显示左侧项目目录，避免把全平台导航误解成项目分支。",
        "首页主按钮可进入P4-T2“5G网络优化结果验证”完整任务。",
    ])

    add_heading(doc, "步骤2：查看项目页和项目四路径", 2)
    add_para(doc, "项目页用于理解整本教材如何从项目进入任务。当前重点观察项目二“5G网络测试”和项目四“5G端到端网络优化”。")
    add_image(doc, images["project"], "图4 项目页：左侧切换项目，主体展示项目链和重点学习路径。", 6.35)
    add_image(doc, images["project_p4"], "图5 项目四路径：从优化实施、结果验证到报告输出形成项目级闭环。", 6.35)
    add_bullets(doc, [
        "项目二承担测试数据和证据形成，项目四承担优化实施、结果验证和报告输出。",
        "P4-T2是当前完整深样章，其他任务主要用于呈现项目闭环和前后承接。",
        "专家可重点判断项目链是否符合5G网优真实工作过程。",
    ])

    add_heading(doc, "步骤3：查看课程能力图谱", 2)
    add_para(doc, "图谱页用于说明课程主链、项目路径、能力节点、资源卡片和评价产出之间的承接关系。当前图谱页已取消右侧上下文栏，优先保证图谱主体的可读面积。")
    add_image(doc, images["diagram_graph"], "图6 课程能力图谱关系模型：图谱不是目录，而是资源、活动和评价的组织暗线。", 6.35)
    add_image(doc, images["graph_top"], "图7 图谱页上半部分：课程主链和重点项目路径。", 6.35)
    add_image(doc, images["graph_resource"], "图8 图谱页资源挂接部分：能力节点与资源卡片、评价产出形成映射。", 6.35)
    add_bullets(doc, [
        "第一层是课程主链，说明学习方向和职业工作过程。",
        "第二层是重点项目路径，说明项目之间如何承接。",
        "第三层是能力节点，说明学生要训练的具体动作。",
        "第四层是资源卡片与评价产出，说明资源为什么出现、服务哪个节点、形成什么产出。",
    ])

    add_heading(doc, "步骤4：查看教师页", 2)
    add_para(doc, "教师页面向任务组织、讲评与复核，不是学生学习正文。专家可判断AI预生成建议、教师审核确认和课堂复核提示是否能支撑真实授课。")
    add_image(doc, images["teacher"], "图9 整书教师页：围绕任务组织、课堂推进、讲评反馈和复核验收组织教师工作。", 6.35)

    add_heading(doc, "步骤5：查看平台支持", 2)
    add_para(doc, "平台支持不是学生学习主界面，而是用于说明任务资源映射、素材子平台边界和交付方式。专家可据此判断资源治理和平台交付链条是否清楚。")
    add_image(doc, images["support_resources"], "图10 平台支持：任务资源映射，说明资源卡片如何服务能力节点、学习活动和评价产出。", 6.35)
    add_image(doc, images["support_material"], "图11 平台支持：素材子平台说明，说明素材入库、媒体审查、教学化加工和资源绑定边界。", 6.35)
    add_image(doc, images["support_delivery"], "图12 平台支持：交付说明，说明资源包输出和直接呈现挂接的差异。", 6.35)

    add_heading(doc, "五、进入P4-T2任务级深样章", 1)
    add_para(doc, "P4-T2“5G网络优化结果验证”是当前完整任务级样章，用于验证学生端、教师端和任务资源是否能形成一节课或两课时的学习闭环。")
    add_heading(doc, "学生学习页：从问题进入、完成活动、形成产出", 2)
    add_image(doc, images["task_student"], "图13 P4-T2学生学习页：学生从任务入口进入，按案例完成学习活动并形成可检查产出。", 6.35)
    add_bullets(doc, [
        "学生页区分课堂带学和自学跟练。",
        "页面以案例和活动组织学习，不只是展示标题或资源入口。",
        "学生需要完成分类、排序、指标判断、依据选择、结论拼装和修正表达等动作。",
    ])

    add_heading(doc, "任务组织页：AI建议、教师确认、讲评和复核", 2)
    add_image(doc, images["task_teacher"], "图14 P4-T2任务组织页上半部分：课堂组织、板书投屏、提问脚本和节奏安排。", 6.35)
    add_image(doc, images["task_teacher_data"], "图15 P4-T2任务组织页下半部分：AI建议、模拟学情、讲评优先级和专业复核提示。", 6.35)
    add_bullets(doc, [
        "教师页保留两课时课堂节奏，学生页不显示课时安排。",
        "模拟学情用于说明未来真实数据回收后的讲评方式，目前不是实际学生数据。",
        "专业复核提示用于提醒专家关注指标阈值、结论口径和案例数据是否正确。",
    ])

    add_heading(doc, "任务资源页：资源转化和媒体治理", 2)
    add_image(doc, images["task_resources"], "图16 P4-T2任务资源页：展示教材资源如何转为学习单、表格、互动、图示和后续可扩展资源。", 6.35)
    add_bullets(doc, [
        "资源不是附件，应服务具体能力节点训练。",
        "原教材中的软件截图、仿真界面和图片需要逐项评估，必要时重绘或替换。",
        "后续可把适合的资源转为动画、互动、小游戏或AI助手支持，但前提是教学目标明确。",
    ])

    doc.add_page_break()
    add_heading(doc, "六、专家重点评审问题", 1)
    add_para(doc, "专家反馈建议尽量具体到页面、节点、资源或任务步骤。下表给出建议评审维度。")
    add_review_table(doc)

    doc.add_page_break()
    add_heading(doc, "七、建议专家填写的反馈格式", 1)
    add_para(doc, "为便于后续迭代，建议每条意见至少包含“问题位置、问题说明、影响、修改建议、优先级”。")
    add_kv_table(doc, [
        ("问题位置", "例如：图谱页第三层P4T2-N03、P4-T2学生学习页案例2、平台支持中的任务资源映射、P4-T2任务资源页。"),
        ("问题说明", "用一句话说明当前哪里不清楚、不正确、不适合教学或不利于平台交付。"),
        ("影响判断", "说明会影响学生自学、教师授课、专业准确性、媒体合规、平台发布或数据安全中的哪一类。"),
        ("修改建议", "尽量给出可执行建议，例如补充术语解释、调整指标阈值、重绘截图、增加对照案例、改资源形态。"),
        ("优先级", "高：不改会影响正确性或教学闭环；中：影响理解和使用效率；低：主要是表达或视觉优化。"),
    ])
    add_callout(
        doc,
        "最需要专家判断的事项",
        "一是5G网优专业内容是否正确；二是学生是否能按页面完成学习闭环；三是教师是否能直接组织教学；四是课程能力图谱是否真正支撑资源和评价；五是媒体、数据和平台交付是否存在正式发布风险。",
        SOFT_GOLD,
    )

    add_heading(doc, "八、当前版本不应被误解为已完成的事项", 1)
    add_bullets(doc, [
        "不代表整本教材所有任务都已完成任务级深样章。",
        "不代表课程能力图谱已经通过通信专业专家最终复核。",
        "不代表原教材全部图片、软件截图、仿真界面都已完成版权和媒体重绘审查。",
        "不代表已接入真实学习数据、真实AI助手、出版社平台接口或正式资源包封装。",
        "不代表已经完成真实一线教师和学生试用。",
    ])

    doc.add_page_break()
    add_heading(doc, "附录：快速操作清单", 1)
    add_step_table(doc, [
        "打开整书预览入口。",
        "在课程页确认整书项目链和P4-T2入口。",
        "切换到项目页，查看项目二和项目四的学习路径。",
        "切换到图谱页，查看课程主链、项目路径、能力节点和资源卡片。",
        "点击“进入P4-T2优化结果验证”，模拟学生学习。",
        "切换P4-T2页面的“任务组织”和“任务资源”，观察AI辅助教师组织和资源转化。",
        "返回整书首页，打开“平台支持”，查看任务资源映射、素材子平台说明和交付说明。",
        "按评审问题表提交改进意见。",
    ])

    doc.save(OUT_DOCX)
    SOURCE_MD.write_text(markdown_source(), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    images = prepare_doc_images()
    build_doc(images)
    print(json.dumps({
        "docx": str(OUT_DOCX),
        "source": str(SOURCE_MD),
        "figures": {k: str(v) for k, v in images.items()},
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
